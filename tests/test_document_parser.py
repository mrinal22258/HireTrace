import os
import tempfile
import pytest
from pathlib import Path
from agents.document_parser import (
    extract_text,
    extract_text_from_bytes,
    infer_document_type,
    compute_file_hash,
)


def test_plain_text_and_markdown_extraction():
    with tempfile.TemporaryDirectory() as tmpdir:
        txt_path = os.path.join(tmpdir, "cv.txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write("# Jane Doe\nSenior Backend Engineer\nPython, Kafka, Distributed Systems.")

        md_path = os.path.join(tmpdir, "notes.md")
        with open(md_path, "w", encoding="utf-8") as f:
            f.write("## Interview Debrief\nStrong grasp of event-driven architectures.")

        txt_extracted = extract_text(txt_path)
        assert "Jane Doe" in txt_extracted
        assert "Distributed Systems" in txt_extracted

        md_extracted = extract_text(md_path)
        assert "Interview Debrief" in md_extracted
        assert "event-driven" in md_extracted


def test_docx_extraction():
    from docx import Document
    with tempfile.TemporaryDirectory() as tmpdir:
        doc_path = os.path.join(tmpdir, "resume.docx")
        doc = Document()
        doc.add_heading("John Smith", level=1)
        doc.add_paragraph("Staff Distributed Systems Engineer with 10 years experience.")
        
        # Add a table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Skill"
        table.cell(0, 1).text = "Level"
        table.cell(1, 0).text = "Python"
        table.cell(1, 1).text = "Expert"
        
        doc.save(doc_path)

        extracted = extract_text(doc_path)
        assert "John Smith" in extracted
        assert "Staff Distributed Systems Engineer" in extracted
        assert "Python | Expert" in extracted or "Python" in extracted


def test_pdf_extraction():
    import pypdf
    with tempfile.TemporaryDirectory() as tmpdir:
        pdf_path = os.path.join(tmpdir, "applicant_cv.pdf")
        writer = pypdf.PdfWriter()
        # Create a blank page with text annotation or draw text
        page = writer.add_blank_page(width=300, height=300)
        
        # We can write PDF content stream
        with open(pdf_path, "wb") as f:
            writer.write(f)

        # Even with an empty/blank PDF, extract_text should not crash
        extracted = extract_text(pdf_path)
        assert isinstance(extracted, str)


def test_legacy_doc_rejection():
    with tempfile.TemporaryDirectory() as tmpdir:
        doc_path = os.path.join(tmpdir, "old_resume.doc")
        with open(doc_path, "wb") as f:
            f.write(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1")  # OLE header
        
        with pytest.raises(ValueError) as exc:
            extract_text(doc_path)
        assert "Legacy .doc format is not supported" in str(exc.value)


def test_unsupported_format_rejection():
    with tempfile.TemporaryDirectory() as tmpdir:
        bad_path = os.path.join(tmpdir, "image.png")
        with open(bad_path, "wb") as f:
            f.write(b"fake image data")

        with pytest.raises(ValueError) as exc:
            extract_text(bad_path)
        assert "Unsupported file format" in str(exc.value)


def test_file_not_found():
    with pytest.raises(FileNotFoundError):
        extract_text("non_existent_file_path_12345.pdf")


def test_infer_document_type():
    assert infer_document_type("jane_doe_cv.pdf") == "cv"
    assert infer_document_type("resume_v2.docx") == "cv"
    assert infer_document_type("technical_interview_transcript.txt") == "interview"
    assert infer_document_type("coding_assessment_results.md") == "assessment"
    assert infer_document_type("system_architecture_rfc.pdf") == "project"


def test_compute_file_hash():
    data = b"Hello HireTrace Pipeline"
    h1 = compute_file_hash(data)
    h2 = compute_file_hash(data)
    assert h1 == h2
    assert len(h1) == 64  # SHA-256 hex length


def test_extract_text_from_bytes():
    txt_bytes = "Software Engineer\n5 years experience".encode("utf-8")
    extracted = extract_text_from_bytes(txt_bytes, "profile.txt")
    assert "Software Engineer" in extracted
