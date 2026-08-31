import os
import io
import time
import json
import socket
import shutil
import tempfile
import threading
import pytest
import requests
from docx import Document

from ui.server import ReusableHTTPServer, HireTraceHandler, root_dir, CASES_DIR, CACHE_DIR, UPLOADS_DIR


def get_free_port():
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
        s.bind(("", 0))
        return s.getsockname()[1]


@pytest.fixture(scope="module")
def server_url():
    port = get_free_port()
    server = ReusableHTTPServer(("127.0.0.1", port), HireTraceHandler)
    t = threading.Thread(target=server.serve_forever, daemon=True)
    t.start()
    time.sleep(0.5)
    return f"http://127.0.0.1:{port}"


def test_upload_single_candidate_multipart(server_url):
    # 1. Prepare a DOCX file for CV
    doc_buffer = io.BytesIO()
    doc = Document()
    doc.add_heading("Marcus Vance", level=1)
    doc.add_paragraph("Principal Distributed Systems Engineer with 8 years experience building Kafka streaming architectures in Python and asyncio.")
    doc.save(doc_buffer)
    doc_buffer.seek(0)

    # 2. Prepare text assessment
    assessment_content = "Evaluated on Raft consensus implementation: 95/100. Flawless concurrency control."

    # 3. Post multipart
    files = {
        "cv_file": ("marcus_resume.docx", doc_buffer.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
        "assessment_file": ("test_eval.txt", assessment_content.encode("utf-8"), "text/plain"),
    }
    data = {
        "name": "Marcus Vance",
        "target_role": "Principal Systems Engineer",
        "interview_notes": "Marcus answered in-depth questions on partition rebalancing and consumer lag metrics.",
        "project_rfc": "Author of RFC 402: Low-Latency Distributed Cache."
    }

    res = requests.post(f"{server_url}/api/candidate/upload?sync=true", data=data, files=files, timeout=300)
    assert res.status_code == 200, f"Failed with {res.status_code}: {res.text}"

    resp_json = res.json()

    assert resp_json.get("success") is True
    cid = resp_json.get("candidate_id", "")
    assert "custom_marcus_vance" in cid
    assert "report" in resp_json
    assert "baseline_a" in resp_json
    assert "raw_documents" in resp_json
    assert "cv" in resp_json["raw_documents"]
    assert resp_json["raw_documents"]["cv"]["filename"] == "marcus_resume.docx"

    try:
        # Check that candidate is in /api/cases
        list_res = requests.get(f"{server_url}/api/cases")
        assert list_res.status_code == 200
        cases_list = list_res.json()
        assert any(c["candidate_id"] == cid for c in cases_list)

        # Check full candidate profile
        full_res = requests.get(f"{server_url}/api/case/{cid}/full")
        assert full_res.status_code == 200
        full_data = full_res.json()
        assert "Marcus Vance" in full_data["documents"]["cv"]
        assert "Kafka" in full_data["documents"]["cv"]
        assert "partition rebalancing" in full_data["documents"]["interview"]
        assert "raw_documents" in full_data

        # Check disk persistence in uploads/
        cand_upload_folder = os.path.join(UPLOADS_DIR, cid)
        assert os.path.exists(cand_upload_folder)
        assert os.path.exists(resp_json["raw_documents"]["cv"]["disk_path"])

    finally:
        # Clean up test artifacts
        case_file = os.path.join(CASES_DIR, f"{cid}.json")
        traj_file = os.path.join(CACHE_DIR, f"{cid}_trajectory.json")
        cand_upload_folder = os.path.join(UPLOADS_DIR, cid)
        for p in (case_file, traj_file):
            if os.path.exists(p):
                try:
                    os.remove(p)
                except Exception:
                    pass
        if os.path.exists(cand_upload_folder):
            try:
                shutil.rmtree(cand_upload_folder)
            except Exception:
                pass


def test_upload_missing_name_error(server_url):
    res = requests.post(f"{server_url}/api/candidate/upload", data={}, files={"cv_file": ("cv.txt", b"content", "text/plain")})
    assert res.status_code == 400
    assert "Missing required field: 'name'" in res.text


def test_upload_missing_cv_error(server_url):
    res = requests.post(f"{server_url}/api/candidate/upload", data={"name": "No CV Applicant"})
    assert res.status_code == 400
    assert "Missing required document" in res.text


def test_upload_legacy_doc_rejection(server_url):
    files = {
        "cv_file": ("old_resume.doc", b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1", "application/msword")
    }
    data = {"name": "Old Word User"}
    res = requests.post(f"{server_url}/api/candidate/upload", data=data, files=files)
    assert res.status_code == 400
    assert "Legacy .doc format is not supported" in res.text
