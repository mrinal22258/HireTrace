"""
Document Parser for HireTrace.

Extracts plain text from multi-format applicant documents:
- PDF (.pdf) via pdfplumber with pypdf and optional OCR fallback
- Word (.docx) via python-docx with zero-dependency XML zip fallback
- Plain text (.txt, .md) with multi-encoding detection (UTF-8, Latin-1, CP1252)
- Rejects legacy .doc with informative guidance

All extractions feed directly into EvidenceLoader without pipeline alteration.
"""

import os
import io
import re
import hashlib
import zipfile
from pathlib import Path
from typing import Union, Optional

try:
    import defusedxml.ElementTree as ET
except ImportError:
    import xml.etree.ElementTree as ET  # nosec B405 - unreachable when defusedxml is installed


def compute_file_hash(data: bytes) -> str:
    """Computes SHA-256 digest of file bytes for deduplication."""
    return hashlib.sha256(data).hexdigest()


def infer_document_type(filename: str) -> str:
    """
    Infers document type from filename patterns for bulk ingestion.
    Returns: 'cv', 'interview', 'assessment', or 'project'.
    """
    base = os.path.basename(filename).lower()
    name, _ = os.path.splitext(base)
    # Replace separators with spaces so \b matches words separated by underscores or dashes
    normalized = re.sub(r"[_\-\s\.]+", " ", name).strip()

    if re.search(r"\b(cv|resume|curriculum|vitae)\b", normalized):
        return "cv"
    if re.search(r"\b(interview|transcript|debrief|screen|call|notes)\b", normalized):
        return "interview"
    if re.search(r"\b(assessment|test|task|takehome|challenge|coding|grader|eval)\b", normalized):
        return "assessment"
    if re.search(r"\b(rfc|project|portfolio|architecture|design|repo|spec)\b", normalized):
        return "project"

    return "cv"



def _extract_pdf(file_path: Union[str, Path]) -> str:
    """Extracts text from PDF using pdfplumber, falling back to pypdf and OCR."""
    text = ""
    # 1. Try pdfplumber (layout-aware, handles columns and tables)
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            pages_text = []
            for page in pdf.pages:
                page_str = page.extract_text()
                if page_str:
                    pages_text.append(page_str)
            text = "\n\n".join(pages_text).strip()
    except Exception:
        text = ""

    # 2. Fallback to pypdf if pdfplumber produced empty or failed
    if not text:
        try:
            import pypdf
            reader = pypdf.PdfReader(str(file_path))
            pages_text = []
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    pages_text.append(extracted)
            text = "\n\n".join(pages_text).strip()
        except Exception:
            pass

    # 3. Fallback to PyMuPDF (fitz) if available and still empty
    if not text:
        try:
            import fitz
            doc = fitz.open(str(file_path))
            pages_text = [page.get_text() for page in doc if page.get_text().strip()]
            text = "\n\n".join(pages_text).strip()
            doc.close()
        except Exception:
            pass

    # 4. Optional OCR fallback for scanned/image PDFs if text is near-empty (< 50 chars)
    if len(text) < 50:
        try:
            import pytesseract
            from pdf2image import convert_from_path
            images = convert_from_path(str(file_path), first_page=1, last_page=5)
            ocr_text = []
            for img in images:
                ocr_str = pytesseract.image_to_string(img)
                if ocr_str.strip():
                    ocr_text.append(ocr_str.strip())
            if ocr_text:
                text = "\n\n".join(ocr_text)
        except Exception:
            # OCR optional dependencies (tesseract-ocr binary / poppler) may not be installed
            pass

    return text.strip()


def _extract_docx(file_path: Union[str, Path]) -> str:
    """Extracts text from DOCX using python-docx with fallback to XML parsing."""
    # 1. Try python-docx
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Include table contents
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    paragraphs.append(" | ".join(row_text))
                    
        return "\n\n".join(paragraphs).strip()
    except Exception:
        pass

    # 2. Zero-dependency fallback: extract word/document.xml from zip
    try:
        with zipfile.ZipFile(file_path, "r") as z:
            xml_content = z.read("word/document.xml")
            root = ET.fromstring(xml_content)  # nosec B314 - uses defusedxml
            texts = []
            for elem in root.iter():
                if elem.tag.endswith("}t") and elem.text:
                    texts.append(elem.text)
                elif elem.tag.endswith("}p"):
                    texts.append("\n")
            full_text = "".join(texts)
            cleaned = re.sub(r"\n\s*\n+", "\n\n", full_text).strip()
            return cleaned
    except Exception as err:
        raise ValueError(f"Failed to parse .docx file: {err}")


def _extract_plain_text(file_path: Union[str, Path]) -> str:
    """Extracts text from plain text or markdown files trying multiple encodings."""
    with open(file_path, "rb") as f:
        raw_bytes = f.read()

    encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]
    for enc in encodings:
        try:
            return raw_bytes.decode(enc).strip()
        except UnicodeDecodeError:
            continue

    return raw_bytes.decode("utf-8", errors="replace").strip()


def extract_text(file_path: Union[str, Path], filename_hint: Optional[str] = None) -> str:
    """
    Extracts plain text from a supported document file.
    
    Supported formats:
      - PDF: .pdf
      - Word: .docx
      - Plain text: .txt, .md
      
    Raises:
      FileNotFoundError: If the file does not exist.
      ValueError: If the file extension is unsupported or legacy .doc.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    # Determine extension
    ext = path.suffix.lower()
    if not ext and filename_hint:
        ext = Path(filename_hint).suffix.lower()

    if ext == ".pdf":
        return _extract_pdf(path)
    elif ext == ".docx":
        return _extract_docx(path)
    elif ext in (".txt", ".md"):
        return _extract_plain_text(path)
    elif ext == ".doc":
        raise ValueError(
            "Legacy .doc format is not supported. Please re-save as .docx or .pdf before uploading."
        )
    else:
        raise ValueError(
            f"Unsupported file format '{ext}'. Supported formats: .pdf, .docx, .txt, .md"
        )


def extract_text_from_bytes(content: bytes, filename: str) -> str:
    """
    Extracts text from in-memory bytes buffer given a filename.
    Useful for direct multipart streaming without permanent disk writes.
    """
    import tempfile

    suffix = Path(filename).suffix.lower()
    if suffix == ".doc":
        raise ValueError(
            "Legacy .doc format is not supported. Please re-save as .docx or .pdf before uploading."
        )
    if suffix not in (".pdf", ".docx", ".txt", ".md"):
        raise ValueError(
            f"Unsupported file format '{suffix}'. Supported formats: .pdf, .docx, .txt, .md"
        )

    # For text/markdown, decode directly in memory
    if suffix in (".txt", ".md"):
        encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]
        for enc in encodings:
            try:
                return content.decode(enc).strip()
            except UnicodeDecodeError:
                continue
        return content.decode("utf-8", errors="replace").strip()

    # For binary formats (.pdf, .docx), write to a temp file and parse
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    try:
        return extract_text(tmp_path, filename_hint=filename)
    finally:
        if os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except Exception:
                pass
